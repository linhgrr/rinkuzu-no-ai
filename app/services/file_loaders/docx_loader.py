"""
DOCX file loader for RAG pipeline
"""
from typing import List, BinaryIO, Dict, Any
from docx import Document
from loguru import logger

from app.core.interfaces.file_loader import FileLoader, DocumentChunk


class DOCXLoader(FileLoader):
    """
    DOCX file loader that extracts text content from Word documents
    
    Features:
    - Text extraction from paragraphs and tables
    - Metadata preservation (paragraph index, filename)
    - Chunking support for large documents
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        """
        Initialize DOCX loader
        
        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_extensions = ["docx", "doc"]
    
    async def load(self, file_stream: BinaryIO, filename: str) -> List[DocumentChunk]:
        """
        Load and process DOCX file into chunks
        
        Args:
            file_stream: DOCX file binary stream
            filename: Original filename for metadata
            
        Returns:
            List of document chunks with content and metadata
        """
        try:
            # Read DOCX content
            doc = Document(file_stream)
            
            logger.info(f"📄 Processing DOCX: {filename}")
            
            all_chunks = []
            full_text = ""
            paragraph_count = 0
            
            # Extract text from paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                
                if para_text:
                    paragraph_count += 1
                    full_text += para_text + "\n\n"
                    
                    # Create paragraph-level chunk if it's reasonably sized
                    if 50 <= len(para_text) <= self.chunk_size:
                        chunk = DocumentChunk(
                            content=para_text,
                            metadata={
                                "filename": filename,
                                "paragraph_index": para_idx,
                                "total_paragraphs": len(doc.paragraphs),
                                "source_type": "docx",
                                "chunk_type": "paragraph",
                                "has_formatting": bool(paragraph.runs)
                            },
                            chunk_index=len(all_chunks)
                        )
                        all_chunks.append(chunk)
            
            # Extract text from tables
            for table_idx, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                
                if table_text.strip():
                    full_text += f"\n--- Table {table_idx + 1} ---\n{table_text}\n\n"
                    
                    # Create table chunk
                    chunk = DocumentChunk(
                        content=table_text.strip(),
                        metadata={
                            "filename": filename,
                            "table_index": table_idx,
                            "total_tables": len(doc.tables),
                            "source_type": "docx",
                            "chunk_type": "table"
                        },
                        chunk_index=len(all_chunks)
                    )
                    all_chunks.append(chunk)
            
            # Create overlapping chunks from full text if needed
            if len(full_text) > self.chunk_size:
                text_chunks = self._create_text_chunks(full_text)
                
                for i, chunk_text in enumerate(text_chunks):
                    chunk = DocumentChunk(
                        content=chunk_text.strip(),
                        metadata={
                            "filename": filename,
                            "total_paragraphs": paragraph_count,
                            "source_type": "docx",
                            "chunk_type": "text_chunk",
                            "chunk_size": len(chunk_text)
                        },
                        chunk_index=len(all_chunks) + i
                    )
                    all_chunks.append(chunk)
            
            logger.info(f"✅ DOCX processed: {filename} -> {len(all_chunks)} chunks")
            return all_chunks
            
        except Exception as e:
            logger.error(f"❌ Failed to process DOCX {filename}: {e}")
            raise ValueError(f"Failed to process DOCX file: {e}")
    
    def supports_file_type(self, file_extension: str) -> bool:
        """Check if this loader supports the given file extension"""
        return file_extension.lower() in self.supported_extensions
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return self.supported_extensions.copy()
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table
        
        Args:
            table: Python-docx table object
            
        Returns:
            Formatted table text
        """
        table_text = ""
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text.append(cell_text)
            
            if any(cell for cell in row_text):  # Skip empty rows
                table_text += " | ".join(row_text) + "\n"
        
        return table_text
    
    def _create_text_chunks(self, text: str) -> List[str]:
        """
        Create overlapping text chunks from large text
        
        Args:
            text: Full text to chunk
            
        Returns:
            List of text chunks with overlap
        """
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at paragraph boundary first
            if end < len(text):
                para_break = text.rfind('\n\n', start, end)
                if para_break > start + self.chunk_size // 3:
                    end = para_break
                else:
                    # Fall back to sentence boundary
                    sentence_break = text.rfind('. ', start, end)
                    if sentence_break > start + self.chunk_size // 2:
                        end = sentence_break + 1
                    else:
                        # Fall back to word boundary
                        space_pos = text.rfind(' ', start, end)
                        if space_pos > start + self.chunk_size // 2:
                            end = space_pos
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = max(start + 1, end - self.chunk_overlap)
            
            # Prevent infinite loop
            if start >= len(text):
                break
        
        return chunks 